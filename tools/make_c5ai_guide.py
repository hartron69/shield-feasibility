"""
Generate C5AI_PCC_GUIDE.docx – a user-facing document explaining how
C5AI+ and the PCC Feasibility Tool work individually and together.

Usage:  python tools/make_c5ai_guide.py
Output: C5AI_PCC_GUIDE.docx  (repo root)
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1A, 0x37, 0x5E)
TEAL        = RGBColor(0x00, 0x6E, 0x7F)
GREEN       = RGBColor(0x1B, 0x5E, 0x20)
AMBER       = RGBColor(0x7E, 0x45, 0x00)
LIGHT_BLUE  = RGBColor(0xD6, 0xEB, 0xF2)
LIGHT_GREEN = RGBColor(0xD9, 0xF0, 0xDA)
LIGHT_AMBER = RGBColor(0xFD, 0xF0, 0xD5)
MID_GREY    = RGBColor(0x44, 0x44, 0x44)
DARK_GREY   = RGBColor(0x22, 0x22, 0x22)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF4, 0xF4, 0xF4)


def _hex(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def _shd(element, fill: RGBColor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(fill))
    element.append(shd)


def set_cell_bg(cell, colour: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    _shd(tcPr, colour)


def add_table_borders(table, colour="CCCCCC"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    b = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), colour)
        b.append(el)
    tblPr.append(b)


def para_shading(para, fill: RGBColor):
    pPr = para._p.get_or_add_pPr()
    _shd(pPr, fill)


class Doc:
    def __init__(self):
        self.doc = Document()
        self._setup()
        self._cover()

    # ── Setup ──────────────────────────────────────────────────────────────────

    def _setup(self):
        for s in self.doc.sections:
            s.top_margin = s.bottom_margin = Cm(2.5)
            s.left_margin = Cm(3.0)
            s.right_margin = Cm(2.5)

        n = self.doc.styles["Normal"]
        n.font.name = "Calibri"
        n.font.size = Pt(10.5)
        n.font.color.rgb = MID_GREY
        n.paragraph_format.space_after = Pt(5)

        for lvl, size, colour, bold in [
            (1, 22, NAVY,  True),
            (2, 15, TEAL,  True),
            (3, 12, NAVY,  True),
            (4, 11, TEAL,  True),
        ]:
            s = self.doc.styles[f"Heading {lvl}"]
            s.font.name = "Calibri"
            s.font.size = Pt(size)
            s.font.bold = bold
            s.font.color.rgb = colour
            s.paragraph_format.space_before = Pt({1:18,2:14,3:10,4:8}[lvl])
            s.paragraph_format.space_after  = Pt({1:6, 2:4, 3:3, 4:2}[lvl])
            if lvl == 4:
                s.font.italic = True

    # ── Cover ──────────────────────────────────────────────────────────────────

    def _cover(self):
        d = self.doc
        for _ in range(3):
            d.add_paragraph()

        # Logo-tekst
        logo = d.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = logo.add_run("Shield Risk Platform")
        r.font.name = "Calibri"
        r.font.size = Pt(34)
        r.font.bold = True
        r.font.color.rgb = NAVY

        sub = d.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = sub.add_run("C5AI+ og PCC Feasibility Tool")
        r2.font.name = "Calibri"
        r2.font.size = Pt(20)
        r2.font.color.rgb = TEAL

        d.add_paragraph()

        tagline = d.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = tagline.add_run("Slik fungerer systemene – hver for seg og sammen")
        r3.font.name = "Calibri"
        r3.font.size = Pt(13)
        r3.font.italic = True
        r3.font.color.rgb = MID_GREY

        d.add_paragraph()
        d.add_paragraph()

        for lbl, val in [
            ("Utarbeidet av:", "Shield Risk Consulting"),
            ("Klassifisering:", "Konfidensielt — Kun for operatør og rådgiver"),
            ("Versjon:", "1.0 — C5AI+ v5.0 / PCC Tool v2.1"),
        ]:
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rl = p.add_run(f"{lbl}  ")
            rl.font.bold = True
            rl.font.size = Pt(11)
            rl.font.color.rgb = NAVY
            rv = p.add_run(val)
            rv.font.size = Pt(11)
            rv.font.color.rgb = MID_GREY

        d.add_page_break()

    # ── Primitive helpers ──────────────────────────────────────────────────────

    def h(self, text: str, level: int = 1):
        p = self.doc.add_paragraph(style=f"Heading {level}")
        p.clear()
        st = self.doc.styles[f"Heading {level}"]
        r = p.add_run(text)
        r.font.name  = st.font.name
        r.font.size  = st.font.size
        r.font.bold  = st.font.bold
        r.font.color.rgb = st.font.color.rgb
        if level == 4:
            r.font.italic = True
        return p

    def p(self, text: str = "", colour=None):
        para = self.doc.add_paragraph()
        self._inline(para, text, colour or MID_GREY)
        para.paragraph_format.space_after = Pt(5)
        return para

    def bullet(self, text: str, level: int = 0):
        style = "List Bullet" if level == 0 else "List Bullet 2"
        para = self.doc.add_paragraph(style=style)
        para.clear()
        self._inline(para, text, MID_GREY)
        return para

    def numbered(self, text: str):
        para = self.doc.add_paragraph(style="List Number")
        para.clear()
        self._inline(para, text, MID_GREY)
        return para

    def code(self, text: str):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent  = Cm(0.8)
        para.paragraph_format.right_indent = Cm(0.8)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(6)
        para_shading(para, LIGHT_GREY)
        r = para.add_run(text)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_GREY
        return para

    def note(self, text: str, bg: RGBColor = None, icon: str = "ⓘ"):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.8)
        para.paragraph_format.space_after = Pt(6)
        if bg:
            para_shading(para, bg)
        r = para.add_run(f"{icon}  {text}")
        r.font.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = TEAL
        return para

    def callout(self, title: str, text: str, bg: RGBColor, title_colour: RGBColor):
        """Shaded paragraph block acting as a callout box."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(8)
        para_shading(para, bg)
        rt = para.add_run(f"{title}  ")
        rt.font.bold = True
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = title_colour
        rb = para.add_run(text)
        rb.font.size = Pt(10)
        rb.font.color.rgb = MID_GREY
        return para

    def hr(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "006E7F")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def table(self, headers, rows,
              col_widths=None,
              header_bg: RGBColor = NAVY,
              stripe_bg: RGBColor = LIGHT_BLUE):
        n = len(headers)
        t = self.doc.add_table(rows=1 + len(rows), cols=n)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Table Grid"
        add_table_borders(t)

        hrow = t.rows[0]
        for i, h in enumerate(headers):
            c = hrow.cells[i]
            set_cell_bg(c, header_bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = c.paragraphs[0].add_run(h)
            r.font.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(10)

        for ri, row in enumerate(rows):
            bg = stripe_bg if ri % 2 == 0 else WHITE
            tr = t.rows[ri + 1]
            for ci, cell_text in enumerate(row):
                c = tr.cells[ci]
                set_cell_bg(c, bg)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                pp = c.paragraphs[0]
                pp.clear()
                self._inline(pp, str(cell_text), DARK_GREY)
                for run in pp.runs:
                    run.font.size = Pt(9.5)

        if col_widths:
            for row in t.rows:
                for ci, w in enumerate(col_widths):
                    row.cells[ci].width = Cm(w)

        self.doc.add_paragraph()
        return t

    @staticmethod
    def _inline(para, text: str, default_colour: RGBColor):
        import re
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                r = para.add_run(text[last:m.start()])
                r.font.color.rgb = default_colour
            full = m.group(0)
            if full.startswith("**"):
                r = para.add_run(m.group(2))
                r.font.bold = True
                r.font.color.rgb = default_colour
            elif full.startswith("*"):
                r = para.add_run(m.group(3))
                r.font.italic = True
                r.font.color.rgb = default_colour
            else:
                r = para.add_run(m.group(4))
                r.font.name = "Courier New"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            last = m.end()
        if last < len(text):
            r = para.add_run(text[last:])
            r.font.color.rgb = default_colour

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path: str):
        self.doc.save(path)
        print(f"Saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# Content
# ══════════════════════════════════════════════════════════════════════════════

def build(d: Doc):

    # ─────────────────────────────────────────────────────────────────────────
    # 1. Innledning
    # ─────────────────────────────────────────────────────────────────────────
    d.h("1. Innledning", 1)
    d.p(
        "Shield Risk Platform består av to separate systemer som kan brukes "
        "uavhengig av hverandre, men som gir størst verdi når de brukes "
        "**sammen**:"
    )
    d.table(
        ["System", "Rolle", "Kjerneteknologi"],
        [
            ["**PCC Feasibility Tool**",
             "Avgjør om en Protected Cell Company (PCC) er en lønnsom forsikringsstruktur for operatøren",
             "Monte Carlo-simulering, aktuariske modeller, kapitalberegning"],
            ["**C5AI+**",
             "Produserer biologisk risikoprognose basert på faktiske feltdata fra lokalitetene",
             "RandomForest ML, nettverksmodellering, kausale effektestimater"],
        ],
        col_widths=[4.5, 7.5, 6],
    )
    d.p(
        "Kort sagt: **C5AI+ svarer på «hva er sannsynligheten for biologiske tap?»** "
        "mens **PCC Feasibility Tool svarer på «hva koster det å forsikre seg mot dem?»**"
    )
    d.callout(
        "Nøkkelpoeng:",
        "De to systemene er løst koblet via én JSON-fil. "
        "PCC-verktøyet kjører helt fint uten C5AI+ — men resultatene blir "
        "mer presise når C5AI+-prognosen er tilgjengelig.",
        LIGHT_BLUE, NAVY,
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 2. PCC Feasibility Tool
    # ─────────────────────────────────────────────────────────────────────────
    d.h("2. PCC Feasibility Tool", 1)

    d.h("2.1 Hva verktøyet gjør", 2)
    d.p(
        "PCC Feasibility Tool simulerer tapsdistribusjoner for en havbruksoperatør "
        "og beregner om en **Protected Cell Company (PCC)** er en bedre løsning enn "
        "tradisjonell forsikring. Det sammenligner fire strategier:"
    )
    d.table(
        ["Strategi", "Beskrivelse"],
        [
            ["Full forsikring", "Alle tap dekkes av kommersielle forsikringsselskaper. Høy premie, ingen kapitaleksponering."],
            ["Hybrid", "Operatøren beholder en egenandel; forsikringen dekker det overskytende."],
            ["PCC Captive Cell", "Operatøren finansierer risiko i en egen beskyttet celle. Lavere premiebelastning, men krever egenkapital som sikkerhet."],
            ["Selvforsikring", "Operatøren bærer all risiko selv. Ingen premie, men full eksponering mot worst-case tap."],
        ],
        col_widths=[4.5, 13.5],
    )

    d.h("2.2 Den statistiske kjernemodellen", 2)
    d.p(
        "Tapene modelleres som en **sammensatt Poisson-LogNormal-prosess** "
        "(Compound Poisson-LogNormal). For hvert simuleringsår trekkes:"
    )
    d.bullet("**Antall hendelser** N ~ Poisson(λ)  — frekvensparameter fra historikk og ekspertanslag")
    d.bullet("**Tap per hendelse** X ~ LogNormal(μ, σ)  — alvorlighetsparameter kalibrert mot TIV og historiske tapsstørrelser")
    d.bullet("**Katastrofehendelse** med sannsynlighet p_cat injiseres oppå attrisjonel tap")
    d.p(
        "Motoren kjører typisk **5 000 simuleringer × 5 år** og produserer en (5000 × 5)-matrise "
        "med årstap. Fra denne matrisen beregnes:"
    )
    d.table(
        ["Nøkkeltall", "Forklaring", "Typisk verdi (Nordic Aqua)"],
        [
            ["E[tap]", "Forventet årlig tap — gjennomsnitt over alle scenarier", "NOK 22,6 M"],
            ["VaR 99,5 %", "Tap som bare overskrides 1 gang per 200 år — tilnærmet SCR", "NOK 134 M"],
            ["CV", "Variasjonskoeffisient (std / gjennomsnitt) — mål på uforutsigbarhet", "2,1"],
            ["TCOR 5 år", "Total Cost of Risk over 5 år for hver strategi", "Avhenger av strategi"],
        ],
        col_widths=[4, 8, 6],
    )

    d.h("2.3 Domenekorrelasjonsmatrise", 2)
    d.p(
        "Risiko fordeles på fire domener: **biologisk, miljø, strukturell og operasjonell**. "
        "En 4×4 korrelasjonsmatrise modellerer samvariasjon mellom domenene — "
        "f.eks. er miljøskader og strukturskader sterkt korrelert (ρ = 0,60) fordi "
        "ekstremvær rammer begge samtidig."
    )
    d.table(
        ["Domene", "Typiske hendelser", "Ekspert-korrelasjon mot miljø"],
        [
            ["Biologisk", "HAB, lakselus, sykdom, oksygenstress", "0,40"],
            ["Miljø", "Storm, strøm, isgang, temperaturekstremer", "— (referanse)"],
            ["Strukturell", "Notskade, fortøyningssvikt, fortøyningsbrudd", "0,60"],
            ["Operasjonell", "Menneskelig feil, driftsavbrudd, tyveri", "0,20"],
        ],
        col_widths=[4, 8.5, 5.5],
    )

    d.h("2.4 Egnethetsvurdering — seks kriterier", 2)
    d.p(
        "Etter simuleringen skårer **SuitabilityEngine** operatøren på seks kriterier "
        "og gir én av tre anbefalinger:"
    )
    d.table(
        ["Kriterium", "Hva som måles"],
        [
            ["Tapsstabilitet", "Lav CV gir mer forutsigbar kapitalbinding i PCC"],
            ["Kapitaleffektivitet", "SCR i % av egenkapital — bør være < 40 %"],
            ["Premieeffektivitet", "PCC-totalkostnad versus markedspremie — hard grense ved +15 %"],
            ["Selvforsikringsegnethet", "Kan operatøren overleve et worst-case tap (VaR 99,5 %)?"],
            ["Diversifikasjonspotensial", "Risikospredning mellom lokaliteter"],
            ["Tapskonsentrasjon", "Dominans fra ett domene øker sårbarhet"],
        ],
        col_widths=[5, 13],
    )
    d.table(
        ["Resultat", "Betingelse"],
        [
            ["✓  ANBEFALT", "Alle seks kriterier bestått, composite score ≥ 70"],
            ["⚠  POTENSIELT EGNET", "Minst fire kriterier bestått, moderate kostnadsforhold"],
            ["✗  IKKE ANBEFALT", "Kritisk kriterium feilet (f.eks. PCC-kostnad > 15 % over FI)"],
        ],
        col_widths=[5, 13],
    )

    d.h("2.5 Statisk modell — styrker og begrensninger", 2)
    d.p(
        "Den statiske modellen bruker **samme λ (frekvens) og μ/σ (alvorlighet) "
        "for alle år og alle scenarier**. Parametrene er kalibrert mot historiske "
        "bransjedata og ekspertvurderinger — ikke mot denne operatørens faktiske "
        "biologiske observasjoner."
    )
    d.callout(
        "Begrensning:",
        "Den statiske modellen vet ikke om det akkurat nå er uvanlig høy "
        "havtemperatur ved lokalitetene, om det er påvist HAB i nærliggende "
        "fjorder, eller om lusenivåene er på vei opp. "
        "Det er her C5AI+ kommer inn.",
        LIGHT_AMBER, AMBER,
    )

    d.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 3. C5AI+
    # ─────────────────────────────────────────────────────────────────────────
    d.h("3. C5AI+ — Biologisk risikoprognose", 1)

    d.h("3.1 Hva C5AI+ gjør", 2)
    d.p(
        "C5AI+ er en modulær AI-plattform som **leser faktiske feltdata fra "
        "lokalitetene** og produserer en biologisk risikoprognose for de neste "
        "1–10 årene. Prognosen angir for hvert risikodomene:"
    )
    d.bullet("Sannsynlighet for at en hendelse inntreffer (hendelsesfrekvens)")
    d.bullet("Forventet tap dersom hendelsen inntreffer (alvorlighet × biomasseverdi)")
    d.bullet("Usikkerhetsbånd (P50 og P90)")
    d.p(
        "C5AI+ kjøres **separat, før PCC-verktøyet**, og resultatet skrives til "
        "en JSON-fil som PCC-verktøyet kan lese."
    )

    d.h("3.2 De fire risikomodellene", 2)
    d.table(
        ["Risikodomene", "Modell", "Status", "Nøkkelinput"],
        [
            ["HAB (skadelig algeoppblomstring)",
             "RandomForestClassifier",
             "Aktiv (Fase 1)",
             "Månedlig havtemperatur, historiske HAB-varsler"],
            ["Lakselus",
             "RandomForestRegressor",
             "Aktiv (Fase 1)",
             "Ukentlige luseregistreringer, temperatur"],
            ["Manet",
             "Prior-only",
             "Placeholder (Fase 2)",
             "Konfigurasjonsprior, ingen ML ennå"],
            ["Patogen",
             "Prior-only",
             "Placeholder (Fase 2)",
             "Konfigurasjonsprior, ingen ML ennå"],
        ],
        col_widths=[4.5, 4, 3, 6.5],
    )

    d.h("3.3 HAB-modellen i detalj (eksempel)", 2)
    d.p(
        "HAB-modellen illustrerer prinsippet som gjelder for alle prognosatorer:"
    )
    d.table(
        ["Steg", "Handling"],
        [
            ["1. Features", "Bygger feature-vektor: måned, havtemperatur, sesongindeks, historisk HAB-rate"],
            ["2. Klassifikasjon", "RandomForest predikerer P(HAB-hendelse) per måned"],
            ["3. Aggregering", "Årlig P(≥1 hendelse) = 1 − ∏(1 − P_måned) over 12 måneder"],
            ["4. Tapsfunksjon", "E[tap | HAB] = biomasseverdi × hab_loss_fraction (LogNormal)"],
            ["5. Ubetinget tap", "E[tap] = P(hendelse) × E[tap | hendelse]"],
            ["6. Fallback", "Hvis < 10 observasjoner: bruker konfig-prior justert for temperaturavvik (±30 %)"],
        ],
        col_widths=[2.5, 15.5],
    )

    d.h("3.4 Geografisk smittespredningsnettverk", 2)
    d.p(
        "Lokaliteter innenfor konfigurerbar avstand (standard: 50 km) kobles i et "
        "**networkx-graf**. Nærliggende lokaliteter deler biologisk risiko via en "
        "risikomultiplikator basert på avstandseksponentiell avfall:"
    )
    d.code("Kantsvekt w = e^(−λ × avstand_km)")
    d.code("Risikomultiplikator = 1.0 + min(0.50, Σ w × 0.15)   [maks 1.50]")
    d.p(
        "Praktisk effekt: en lokalitet med tre naboer innenfor 20 km kan få opptil "
        "50 % forhøyet forventet tap — fordi HAB-vannmasser og lus "
        "sprer seg mellom anlegg."
    )

    d.h("3.5 Datakvalitetsflagg", 2)
    d.p(
        "C5AI+ er åpen om usikkerhet. Hvert nettsted og prognose merkes med "
        "ett av fire flagg:"
    )
    d.table(
        ["Flagg", "Datadekning", "Modell som brukes", "Konfidensnivå"],
        [
            ["`SUFFICIENT`", "≥ 70 %, ≥ 24 mnd.", "RandomForest ML", "Høy"],
            ["`LIMITED`",    "40–69 %",            "Blanding prior/ML", "Medium"],
            ["`POOR`",       "10–39 %",            "Prior med historisk justering", "Lav"],
            ["`PRIOR_ONLY`", "< 10 %",             "Ren konfig-prior", "Lav"],
        ],
        col_widths=[3.5, 3.5, 7, 4],
    )
    d.note(
        "Selv med PRIOR_ONLY produserer C5AI+ et gyldig resultat. "
        "Prognosen er da basert på bransjenormer, ikke på operatørens egne data. "
        "Flagget propageres gjennom til PCC-rapportene slik at styret alltid ser "
        "grunnlaget for prognosene."
    )

    d.h("3.6 Valgfri CATE-modul (Fase 3)", 2)
    d.p(
        "`CATEModule` bruker en **T-Learner meta-learner** (econml-biblioteket) for "
        "å estimere den *kausale* effekten av spesifikke tiltak — f.eks.: "
        "«hva er den faktiske tapsreduksjonen av tidlig høsting ved HAB-fare, "
        "isolert fra andre faktorer?» "
        "Denne modulen er ikke aktiv i standard kjøring og krever "
        "`pip install econml`."
    )

    d.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 4. Integrasjon
    # ─────────────────────────────────────────────────────────────────────────
    d.h("4. Integrasjon — slik fungerer de sammen", 1)

    d.h("4.1 Arkitekturoversikt", 2)
    d.code(
        "┌───────────────────────────────────────────────────────────────────┐\n"
        "│  STEG 1: C5AI+ pipeline  (kjøres separat, én gang per sesong)    │\n"
        "│                                                                   │\n"
        "│  Feltdata (temperatur, lus, HAB-varsler, GPS)                     │\n"
        "│       │                                                           │\n"
        "│       ▼                                                           │\n"
        "│  ForecastPipeline.run()                                           │\n"
        "│    ├── DataLoader        → validerer og standardiserer data       │\n"
        "│    ├── SiteRiskNetwork   → geografisk smittespredning             │\n"
        "│    ├── HABForecaster     → P(HAB) per lokalitet per år            │\n"
        "│    ├── LiceForecaster    → forventet lusepåslag                   │\n"
        "│    └── Aggregering       → operator-nivå total E[tap]             │\n"
        "│       │                                                           │\n"
        "│       ▼                                                           │\n"
        "│  risk_forecast.json  ◄── det eneste som overføres til steg 2     │\n"
        "└─────────────────────────────────┬─────────────────────────────────┘\n"
        "                                  │  c5ai_vs_static_ratio\n"
        "                                  │  loss_breakdown_fractions\n"
        "                                  ▼\n"
        "┌───────────────────────────────────────────────────────────────────┐\n"
        "│  STEG 2: PCC Feasibility Tool  (kjøres av rådgiver / GUI)        │\n"
        "│                                                                   │\n"
        "│  sample_input.json  (med c5ai_forecast_path-felt satt)           │\n"
        "│       │                                                           │\n"
        "│       ▼                                                           │\n"
        "│  MonteCarloEngine.run()                                           │\n"
        "│    ├── Statisk simulering  → (N×T)-matrise med årstap            │\n"
        "│    ├── C5AI+-skalering     → matrise × scale_factor              │\n"
        "│    ├── Bio-nedbrytning     → tap fordelt på HAB/lus/manet/path.  │\n"
        "│    └── Domenekorrelas.     → 4×4 kov.matrise på domenebidrag     │\n"
        "│       │                                                           │\n"
        "│       ▼                                                           │\n"
        "│  Strategi-analyse → SCR → Egnethet → PDF-rapport                 │\n"
        "└───────────────────────────────────────────────────────────────────┘"
    )

    d.h("4.2 Hva risk_forecast.json inneholder", 2)
    d.p(
        "PCC-verktøyet bruker bare **to felter** fra den store JSON-filen:"
    )
    d.code(
        '{\n'
        '  "operator_aggregate": {\n'
        '    "total_expected_annual_loss": 14280000,\n'
        '    "c5ai_vs_static_ratio": 0.632,\n'
        '    "loss_breakdown_fractions": {\n'
        '      "hab":       0.5156,\n'
        '      "lice":      0.2533,\n'
        '      "jellyfish": 0.0924,\n'
        '      "pathogen":  0.1387\n'
        '    }\n'
        '  }\n'
        '}'
    )
    d.table(
        ["Felt", "Forklaring", "Effekt i PCC-verktøyet"],
        [
            ["`c5ai_vs_static_ratio`",
             "C5AI+sitt forventede årstap delt på den statiske modellens forventede årstap",
             "Hele (N×T)-tapmatrisen multipliseres med denne skalaren. Ratio < 1 = C5AI+ anslår lavere risiko enn statisk modell; > 1 = høyere."],
            ["`loss_breakdown_fractions`",
             "Andel av biologisk tap per risikodomene (HAB / lus / manet / patogen). Summerer til 1,0.",
             "Distribuerer de skalerte tapene på biologiske underdomener for disaggregert rapportering i PDF og GUI."],
        ],
        col_widths=[4.5, 6, 7.5],
    )

    d.h("4.3 Skaleringen i praksis — et talleksempel", 2)
    d.p(
        "Anta at den statiske Monte Carlo-modellen gir E[tap] = 22,6 M NOK per år. "
        "C5AI+ analyserer dataene og anslår at biologisk risiko dette året er lavere "
        "enn normalt (kjølig sommer, lite lus). Resultatet:"
    )
    d.table(
        ["", "Verdi"],
        [
            ["Statisk E[tap]", "NOK 22,6 M"],
            ["C5AI+ E[tap]", "NOK 14,3 M"],
            ["**c5ai_vs_static_ratio**", "**0,632**"],
            ["Effekt på VaR 99,5 %", "NOK 134 M × 0,632 = NOK 84,7 M"],
            ["Effekt på SCR (PCC)", "Redusert med ~37 % → lavere kapitalkrav"],
            ["Effekt på TCOR", "Lavere forventet tap → PCC mer konkurransedyktig vs. FI"],
        ],
        col_widths=[7, 11],
        stripe_bg=LIGHT_GREEN,
    )
    d.note(
        "Formen på tapsdistribusjonen (halestruktur, korrelasjonsmønster) endres ikke "
        "av skaleringen — bare nivået flyttes. Dette er bevisst design: "
        "Monte Carlo-metodikken forblir aktuarisk korrekt."
    )

    d.h("4.4 Informasjon som IKKE overføres", 2)
    d.p(
        "Det er like viktig å forstå hva som *ikke* overføres mellom systemene:"
    )
    d.table(
        ["C5AI+ vet ikke om...", "PCC-verktøyet vet ikke om..."],
        [
            ["Forsikringsstrategier (PCC, hybrid, FI)", "Spesifikke temperaturer ved lokalitetene"],
            ["SCR-krav og kapitalregler", "Faktiske luseregistreringer per uke"],
            ["Gjenforsikringsstruktur og loading", "GPS-koordinater og nettverk mellom lokaliteter"],
            ["Solvency-ratio og egenkapital", "Nøyaktige biologiske sannsynligheter"],
            ["TCOR-beregninger og strategi-TCOR", "Hvilken dataserie som lå bak ratio-estimatet"],
        ],
        col_widths=[9, 9],
        header_bg=TEAL,
    )

    d.h("4.5 Fallback-logikk", 2)
    d.p(
        "Hvis `risk_forecast.json` mangler, er korrupt, eller `c5ai_forecast_path` "
        "ikke er satt i input-JSON, faller Monte Carlo-motoren stille tilbake til "
        "den statiske modellen:"
    )
    d.code(
        "# Fra models/monte_carlo.py\n"
        "except FileNotFoundError:\n"
        "    warnings.warn('C5AI+ forecast file not found. Falling back to static model.')\n"
        "    return annual_losses, None, None, False   # uendret matrise"
    )
    d.p(
        "Ingen kræsj, ingen feilmelding til brukeren — kun en `RuntimeWarning` i "
        "terminalloggene. `SimulationResults.c5ai_enriched` settes til `False`, "
        "og rapporten merker resultatet som basert på statisk modell."
    )

    d.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 5. Praktisk bruksanvisning
    # ─────────────────────────────────────────────────────────────────────────
    d.h("5. Praktisk bruksanvisning", 1)

    d.h("5.1 Kjøre kun PCC Feasibility Tool (statisk modus)", 2)
    d.p(
        "Ingen C5AI+-data nødvendig. Start backend og frontend som normalt:"
    )
    d.code(
        "# Backend\n"
        "uvicorn backend.main:app --reload --port 8000\n\n"
        "# Frontend\n"
        "cd frontend && npm run dev\n"
        "# Åpne http://localhost:5173"
    )
    d.p(
        "Fyll inn operatørprofilen i GUI-en, klikk **Kjør analyse**. "
        "`c5ai_enriched` i rapporten vil vise `false`."
    )

    d.h("5.2 Kjøre med C5AI+-integrasjon (anbefalt)", 2)
    d.table(
        ["Steg", "Kommando / handling", "Resultat"],
        [
            ["1",
             "Kjør C5AI+-pipeline:\n`python examples/run_c5ai_demo.py`",
             "Produserer `examples/demo_risk_forecast.json`"],
            ["2",
             "Legg filstien inn i operatør-JSON:\n`\"c5ai_forecast_path\": \"examples/demo_risk_forecast.json\"`",
             "PCC-verktøyet vet nå hvor prognosen ligger"],
            ["3",
             "Kjør PCC-verktøyet:\n`python main.py`  eller via GUI",
             "Monte Carlo skalerer tapene med C5AI+-ratio"],
            ["4",
             "Les rapport",
             "Rapporten viser `C5AI+ beriket: true` og biologisk risikofordeling"],
        ],
        col_widths=[1.5, 8.5, 8],
    )

    d.h("5.3 Oppdateringsfrekvens", 2)
    d.p(
        "C5AI+-prognosen bør fornyes når feltdataene endres vesentlig:"
    )
    d.table(
        ["Triggerhendelse", "Anbefalt handling"],
        [
            ["Ny havtemperatursesong (vår/høst)", "Kjør C5AI+ pipeline på nytt"],
            ["Påvist HAB i nærliggende fjord", "Kjør C5AI+ pipeline på nytt — ratio vil øke"],
            ["Lusenivå overskrider behandlingsterskel", "Kjør C5AI+ pipeline på nytt"],
            ["Nytt PCC-solvensregnskapsår", "Kjør PCC-verktøyet på nytt med eksisterende forecast"],
            ["Vesentlig endring i biomasse eller lokaliteter", "Kjør begge systemer på nytt"],
        ],
        col_widths=[8, 10],
    )

    d.h("5.4 Datakrav for C5AI+", 2)
    d.table(
        ["Datakilde", "Minimum for ML", "Ideelt", "Konsekvens hvis mangler"],
        [
            ["Havtemperatur (månedlig)", "24 måneder", "5 år", "Bruker prior med temp-justering (LIMITED/POOR)"],
            ["Saltholdighet", "—", "Anbefalt", "Ingen direkte konsekvens"],
            ["Klorofyll-a", "—", "Nyttig for HAB", "HAB-modell mister viktig feature"],
            ["Luseregistreringer (ukentlig)", "8 uker", "2 år", "Luce-modell faller tilbake til prior"],
            ["HAB-varsler (historisk)", "Ingen krav", "3+ hendelser", "HAB-prior blandes ikke med historikk"],
            ["GPS-koordinater per lokalitet", "Obligatorisk", "—", "Nettverksrisiko-modulen deaktiveres"],
        ],
        col_widths=[4.5, 3, 3, 7.5],
    )

    d.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 6. Filreferanse
    # ─────────────────────────────────────────────────────────────────────────
    d.h("6. Filreferanse", 1)

    d.h("6.1 Nøkkelfiler", 2)
    d.table(
        ["Fil", "Plassering", "Formål"],
        [
            ["`sample_input.json`",
             "`data/`",
             "Statisk operatørprofil (Nordic Aqua). Legg til `c5ai_forecast_path`-felt for å aktivere C5AI+."],
            ["`risk_forecast.json`",
             "Valgfritt — oppgis av bruker",
             "C5AI+-prognose. Produseres av `ForecastPipeline`. Leses av `MonteCarloEngine`."],
            ["`demo_risk_forecast.json`",
             "`examples/`",
             "Ferdig eksempelfil fra demo-kjøringen. Kan brukes direkte for testing."],
            ["`risk_forecast_example.json`",
             "`examples/`",
             "Statisk referansefil som følger med repoet. Brukes i tester."],
            ["`run_c5ai_demo.py`",
             "`examples/`",
             "Kjørbart demoskript: bygger syntetisk C5AI+-input, kjører pipeline, integrerer med PCC-verktøyet."],
            ["`c5ai_settings.py`",
             "`c5ai_plus/config/`",
             "Alle C5AI+-konstanter: terskler, prior-sannsynligheter, nettverksparametere."],
        ],
        col_widths=[5, 3.5, 9.5],
    )

    d.h("6.2 Slik aktiveres C5AI+ i sample_input.json", 2)
    d.p("Legg til ett felt i eksisterende JSON (stien er relativ til der du kjører `python main.py` fra):")
    d.code(
        '{\n'
        '  "name": "Nordic Aqua Partners AS",\n'
        '  ...\n'
        '  "c5ai_forecast_path": "examples/demo_risk_forecast.json"\n'
        '}'
    )
    d.p("Fjern feltet (eller sett det til `null`) for å gå tilbake til statisk modus.")

    d.h("6.3 C5AI+ via Python-API (programmatisk)", 2)
    d.code(
        "from c5ai_plus.pipeline import ForecastPipeline\n"
        "from c5ai_plus.data_models.biological_input import (\n"
        "    C5AIOperatorInput, SiteMetadata, EnvironmentalObservation\n"
        ")\n\n"
        "# Bygg inputobjekt med feltdata\n"
        "operator_input = C5AIOperatorInput(\n"
        "    operator_id='min-operatoer-001',\n"
        "    operator_name='Min Oppdrett AS',\n"
        "    sites=[SiteMetadata(site_id='s1', latitude=60.45, longitude=6.10,\n"
        "                        biomass_tonnes=3000, biomass_value_nok=216_000_000)],\n"
        "    env_observations=[EnvironmentalObservation(\n"
        "        site_id='s1', year=2025, month=7, sea_temp_celsius=15.8\n"
        "    )],\n"
        "    forecast_years=5,\n"
        ")\n\n"
        "# Kjør pipeline\n"
        "pipeline = ForecastPipeline(verbose=True)\n"
        "forecast = pipeline.run(\n"
        "    operator_input,\n"
        "    static_mean_annual_loss=22_600_000,   # fra PCC-verktoyets statiske kjoring\n"
        "    output_path='risk_forecast.json',\n"
        ")\n\n"
        "print(f'Skalafaktor: {forecast.operator_aggregate.c5ai_vs_static_ratio:.3f}')"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 7. Fremtidige utvidelser
    # ─────────────────────────────────────────────────────────────────────────
    d.h("7. Fremtidige utvidelser (veikart)", 1)
    d.table(
        ["Fase", "Komponent", "Innhold", "Status"],
        [
            ["1", "HAB-modell", "RandomForest på temperatur og historiske HAB-varsler", "Aktiv"],
            ["1", "Luse-modell", "RandomForest på ukentlige luseregistreringer", "Aktiv"],
            ["2", "Manet-modell", "ML-modell (krever datasett fra Havforskningsinstituttet)", "Planlagt"],
            ["2", "Patogen-modell", "ML-modell (krever datasett fra Mattilsynet)", "Planlagt"],
            ["2", "HAB-forbedring", "Klorofyll-a som feature — øker presisjonen vesentlig", "Planlagt"],
            ["3", "CATE-modul", "T-Learner (econml) for kausale behandlingseffekter", "Scaffoldet"],
            ["3", "Læringsloop", "Inkrementell retrening; shadow mode med manuell godkjenning", "Design"],
            ["3", "Dataintegrering", "Direkte kobling til NorKyst800, BarentsWatch, Mattilsynet", "Design"],
        ],
        col_widths=[1.5, 4, 10, 2.5],
    )
    d.note(
        "Fase 2 og 3 krever datadelingsavtaler med eksterne parter. "
        "Fase 1 er produksjonsklar og fungerer med operatørens egne "
        "miljøobservasjoner og luseregistreringer."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # Footer
    # ─────────────────────────────────────────────────────────────────────────
    d.doc.add_paragraph()
    d.hr()
    footer = d.doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "Shield Risk Platform — C5AI+ v5.0 / PCC Feasibility Tool v2.1  "
        "© Shield Risk Consulting  —  Konfidensielt. Ikke distribuer uten tillatelse."
    )
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "C5AI_PCC_GUIDE.docx"
    d = Doc()
    build(d)
    d.save(str(out))
