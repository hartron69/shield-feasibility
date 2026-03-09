"""
Convert BRUKERMANUAL_GUI.md to a styled Word document.
Usage: python tools/md_to_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1A, 0x37, 0x5E)   # headings
TEAL        = RGBColor(0x00, 0x6E, 0x7F)   # h2 / accent
LIGHT_BLUE  = RGBColor(0xD6, 0xEB, 0xF2)   # table header bg
MID_GREY    = RGBColor(0x55, 0x55, 0x55)   # body text
CODE_BG     = RGBColor(0xF4, 0xF4, 0xF4)   # code block shading
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY   = RGBColor(0x33, 0x33, 0x33)


def set_cell_bg(cell, colour: RGBColor):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_colour = f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def set_cell_borders(table):
    """Add thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_run_formatted(para, text, bold=False, italic=False, code=False,
                      colour=None, size=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    if colour:
        run.font.color.rgb = colour
    if size:
        run.font.size = Pt(size)
    return run


def parse_inline(para, text, default_colour=None):
    """Render **bold**, *italic*, and `code` inline markup into a paragraph."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            if default_colour:
                r.font.color.rgb = default_colour
        full = m.group(0)
        if full.startswith("**"):
            add_run_formatted(para, m.group(2), bold=True, colour=default_colour)
        elif full.startswith("*"):
            add_run_formatted(para, m.group(3), italic=True, colour=default_colour)
        else:
            add_run_formatted(para, m.group(4), code=True)
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        if default_colour:
            r.font.color.rgb = default_colour


class DocBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._add_cover()

    # ── Document-level setup ──────────────────────────────────────────────────

    def _setup_styles(self):
        doc = self.doc
        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        # Default body style
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = MID_GREY
        style.paragraph_format.space_after = Pt(6)

        # Heading 1
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(22)
        h1.font.bold = True
        h1.font.color.rgb = NAVY
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after  = Pt(6)

        # Heading 2
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(15)
        h2.font.bold = True
        h2.font.color.rgb = TEAL
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after  = Pt(4)

        # Heading 3
        h3 = doc.styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = NAVY
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after  = Pt(3)

        # Heading 4
        h4 = doc.styles["Heading 4"]
        h4.font.name = "Calibri"
        h4.font.size = Pt(11)
        h4.font.bold = True
        h4.font.italic = True
        h4.font.color.rgb = TEAL
        h4.paragraph_format.space_before = Pt(8)
        h4.paragraph_format.space_after  = Pt(2)

    def _add_cover(self):
        doc = self.doc
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("Shield Risk Platform")
        r.font.name = "Calibri"
        r.font.size = Pt(32)
        r.font.bold = True
        r.font.color.rgb = NAVY

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = sub.add_run("Brukermanual – Grafisk grensesnitt (GUI)")
        r2.font.name = "Calibri"
        r2.font.size = Pt(18)
        r2.font.color.rgb = TEAL

        doc.add_paragraph()

        meta_lines = [
            ("Versjon:", "2.1  —  Nettbasert GUI med Pooling-analyse"),
            ("Utarbeidet av:", "Shield Risk Consulting"),
            ("Klassifisering:", "Konfidensielt — Kun for operatør og rådgiver"),
            ("Plattform:", "Nettleser (Chrome / Edge / Firefox) + Python-backend"),
        ]
        for label, value in meta_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_lbl = p.add_run(f"{label} ")
            r_lbl.font.bold = True
            r_lbl.font.color.rgb = NAVY
            r_lbl.font.size = Pt(11)
            r_val = p.add_run(value)
            r_val.font.color.rgb = MID_GREY
            r_val.font.size = Pt(11)

        doc.add_page_break()

    # ── Helpers ───────────────────────────────────────────────────────────────

    def heading(self, text, level):
        style = f"Heading {level}"
        p = self.doc.add_paragraph(style=style)
        p.clear()
        run = p.add_run(text)
        target = self.doc.styles[style]
        run.font.name  = target.font.name
        run.font.size  = target.font.size
        run.font.bold  = target.font.bold
        run.font.color.rgb = target.font.color.rgb
        if level == 4:
            run.font.italic = True
        return p

    def body(self, text):
        p = self.doc.add_paragraph()
        parse_inline(p, text, default_colour=MID_GREY)
        p.paragraph_format.space_after = Pt(5)
        return p

    def bullet(self, text, level=0):
        style = "List Bullet" if level == 0 else "List Bullet 2"
        p = self.doc.add_paragraph(style=style)
        p.clear()
        parse_inline(p, text, default_colour=MID_GREY)
        return p

    def note(self, text):
        """Indented italic note paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"ⓘ  {text}")
        r.font.italic = True
        r.font.color.rgb = TEAL
        r.font.size = Pt(10)
        return p

    def code_block(self, text):
        """Shaded monospace block for commands / JSON."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        # shade via XML
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        pPr.append(shd)
        r = p.add_run(text)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_GREY
        return p

    def table(self, headers, rows, col_widths=None):
        """Add a styled table with a navy header row."""
        n_cols = len(headers)
        t = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Table Grid"
        set_cell_borders(t)

        # Header row
        hdr_row = t.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            set_cell_bg(cell, NAVY)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.font.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(10)

        # Data rows
        for ri, row in enumerate(rows):
            tr = t.rows[ri + 1]
            bg = LIGHT_BLUE if ri % 2 == 0 else WHITE
            for ci, cell_text in enumerate(row):
                cell = tr.cells[ci]
                set_cell_bg(cell, bg)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.clear()
                parse_inline(p, str(cell_text), default_colour=DARK_GREY)
                for run in p.runs:
                    run.font.size = Pt(9.5)

        # Column widths
        if col_widths:
            for ri2, row2 in enumerate(t.rows):
                for ci2, width in enumerate(col_widths):
                    row2.cells[ci2].width = Cm(width)

        self.doc.add_paragraph()  # spacing after table
        return t

    def horizontal_rule(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "006E7F")
        pBdr.append(bottom)
        pPr.append(pBdr)


# ── Content builder ───────────────────────────────────────────────────────────

def build(b: DocBuilder):
    doc = b.doc

    # ── 1. Systemoversikt ─────────────────────────────────────────────────────
    b.heading("1. Systemoversikt", 1)
    b.body(
        "Shield Risk Platform er et beslutningsstøtteverktøy for havbruksoperatører som vurderer "
        "om en **Protected Cell Company (PCC)** er en egnet forsikringsstruktur. Verktøyet "
        "simulerer tapsdistribusjoner med Monte Carlo-metoden og sammenligner fire strategier:"
    )
    b.table(
        ["Strategi", "Beskrivelse"],
        [
            ["Full forsikring", "Tradisjonell kommersiell forsikring dekker alle tap"],
            ["Hybrid", "Kombinasjon av egenrisiko og forsikring"],
            ["PCC Captive Cell", "Operatøren finansierer risiko i en beskyttet celle"],
            ["Selvforsikring", "Operatøren bærer all risiko selv"],
        ],
        col_widths=[5, 12],
    )
    b.body(
        "Dersom Pooling er aktivert, analyseres i tillegg en **Pooled PCC Cell** der operatøren "
        "inngår i et syntetisk risikodelingsbasseng med likeverdige aktører."
    )
    b.heading("Nøkkelbegreper", 3)
    b.table(
        ["Begrep", "Forklaring"],
        [
            ["VaR 99,5 %", "Value at Risk ved 99,5 % konfidensnivå. Tap som bare overskrides i 0,5 % av scenariene. Brukes som tilnærming til SCR."],
            ["CV", "Variasjonskoeffisient (standardavvik / forventningsverdi). Lav CV = jevnt tap; høy CV = stor usikkerhet."],
            ["TCOR", "Total Cost of Risk over 5 år – summen av alle direkte og indirekte kostnader."],
            ["E[tap]", "Forventet årlig tap – gjennomsnittet over alle Monte Carlo-simuleringer."],
            ["SCR", "Solvency Capital Requirement – kapitalkrav; tilnærmet lik VaR 99,5 %."],
        ],
        col_widths=[4, 13],
    )

    # ── 2. Oppstart ───────────────────────────────────────────────────────────
    b.heading("2. Oppstart", 1)
    b.heading("Forutsetninger", 3)
    b.table(
        ["Komponent", "Versjonskrav"],
        [
            ["Python", "3.10 eller nyere, med pakker fra requirements.txt"],
            ["Node.js", "18 eller nyere (kun nødvendig ved første gangs oppsett)"],
            ["Nettleser", "Chrome, Edge eller Firefox – siste versjon anbefalt"],
        ],
        col_widths=[5, 12],
    )
    b.heading("Start backend", 3)
    b.code_block("# Fra prosjektets rotkatalog\nuvicorn backend.main:app --reload --port 8000")
    b.body("Forventet utskrift:")
    b.code_block("INFO:     Uvicorn running on http://127.0.0.1:8000\nINFO:     Application startup complete.")
    b.heading("Start frontend (utviklingsmodus)", 3)
    b.code_block("cd frontend\nnpm install          # bare første gang\nnpm run dev")
    b.body("Åpne nettleseren på **http://localhost:5173**.")
    b.heading("Prod-bygg (valgfritt)", 3)
    b.code_block("cd frontend && npm run build\n# Statiske filer i frontend/dist/ – serveres av backend på port 8000")

    # ── 3. Venstre panel ──────────────────────────────────────────────────────
    b.heading("3. Venstre panel – Inndata", 1)
    b.body(
        "Venstrekolonnen er delt inn i fem trekkspillseksjoner. "
        "Klikk på seksjonstittel for å åpne eller lukke."
    )

    b.heading("3.1 Operatørprofil", 2)
    b.table(
        ["Felt", "Beskrivelse", "Eksempel"],
        [
            ["Operatørnavn", "Navn på selskapet", "Nordic Aqua Partners AS"],
            ["Land", "Driftsnasjon", "Norge"],
            ["Antall lokaliteter", "Totalt antall produksjonslokaliteter", "3"],
            ["Total biomasse (t)", "Total stående biomasse på tvers av alle lokaliteter", "9 200 t"],
        ],
        col_widths=[5, 9, 5],
    )
    b.heading("Biomasseverdsetting", 4)
    b.body("Verktøyet beregner en **foreslått verdi** per tonn basert på formelen:")
    b.code_block("Foreslått verdi = Referansepris (NOK/kg) × 1 000 × Realisasjonsfaktor × (1 – Forsiktighetskutt)")
    b.table(
        ["Felt", "Standard", "Beskrivelse"],
        [
            ["Referansepris (NOK/kg)", "80", "Markedspris for laks"],
            ["Realisasjonsfaktor", "0,90", "Andel av markedspris oppnådd ved nødslakt"],
            ["Forsiktighetskutt", "0,10", "Konservativ buffer for prisusikkerhet"],
            ["Foreslått verdi (NOK/t)", "64 800", "Beregnet automatisk, skrivebeskyttet"],
            ["Anvendt verdi (NOK/t)", "64 800", "Redigerbar – trykk «Bruk foreslått» for å tilbakestille"],
        ],
        col_widths=[6, 3, 9],
    )
    b.note(
        "Hvis du kjenner den faktiske forsikringsverdien per tonn, skriv den direkte inn i "
        "Anvendt verdi. Feltet viser da en oransje «OVERSTYRT»-badge som påminnelse."
    )
    b.body("Følgende felter beregnes automatisk og er skrivebeskyttet (markert med lås-ikon):")
    b.bullet("**Årsomsetning (NOK)** = Biomasse × Anvendt verdi × 1,35")
    b.bullet("**Årlig premie (NOK)** = Årsomsetning × 2,17 %")

    b.heading("3.2 Modellinnstillinger", 2)
    b.table(
        ["Felt", "Standard", "Beskrivelse"],
        [
            ["Antall simuleringer", "5 000", "Monte Carlo-iterasjoner. 5 000 = ~1 sek; 20 000 = ~4 sek."],
            ["Domenekorrelasjon", "expert_default", "Korrelasjonsprofil mellom biologiske, miljø-, strukturelle og operasjonelle tapsdrivere."],
            ["Generer PDF", "Ja", "Produserer en 13-siders styrerapport i PDF-format."],
            ["Bruk historikkalibrering", "Nei", "Kalibrerer Monte Carlo mot historiske tapsposter."],
        ],
        col_widths=[5, 3.5, 9],
    )
    b.heading("Tilgjengelige korrelasjonsprofiler", 4)
    b.table(
        ["Profil", "Beskrivelse"],
        [
            ["`expert_default`", "Ekspertkalibrert: moderat bio–miljø (0,40), høy miljø–struktur (0,60). Anbefalt for norsk oppdrett."],
            ["`uncorrelated`", "Alle domener uavhengige – kun for benchmarking."],
            ["`high_correlation`", "Sterk kopling mellom alle domener; konservativt worst-case."],
        ],
        col_widths=[5, 12],
    )

    b.heading("3.3 Strategiinnstillinger", 2)
    b.table(
        ["Felt", "Beskrivelse"],
        [
            ["Strategi", "Velg primærstrategi. Alle fire strategier beregnes alltid; dette bestemmer hvilken PCC-type som fremheves."],
            ["Egenrisiko (NOK)", "Opptil dette beløpet bærer operatøren tapet selv. Tomt felt = verktøyet beregner optimal egenrisiko."],
        ],
        col_widths=[5, 12],
    )

    b.heading("3.4 Risikoreduserende tiltak", 2)
    b.body(
        "Velg hvilke tiltak operatøren har implementert eller vurderer. "
        "Tiltakene påvirker tapssannsynlighet og -alvorlighet i den mitigerte kjøringen."
    )
    for group, rows in [
        ("Strukturelle tiltak", [
            ["Sterkere not", "Forsterket notmateriale", "35 %", "30 %", "450 000"],
            ["Sterkere fortøyning", "Oppgradert fortøyningssystem", "30 %", "25 %", "300 000"],
            ["Deformasjonsovervåking", "Sanntidssensorer for notgeometri og fortøyningsavvik", "30 %", "20 %", "250 000"],
        ]),
        ("Miljøtiltak", [
            ["Miljøsensorer", "Temperatur, oksygen og strømmåling", "25 %", "20 %", "200 000"],
            ["Maneter-avskjerming", "Fysisk og biologisk beskyttelse", "40 %", "35 %", "180 000"],
        ]),
        ("Operasjonelle tiltak", [
            ["Opplæringsprogram", "Systematisk HMS- og beredskapsopplæring", "20 %", "15 %", "120 000"],
            ["Risikostyrer", "Dedikert risikostyringsansvarlig", "—", "25 %", "800 000"],
            ["Stormkontingensplan", "Dokumentert prosedyre ved ekstremvær", "—", "20 %", "50 000"],
            ["Beredskapsprogram", "Responsprotokoller for kritiske hendelser", "—", "20 %", "80 000"],
        ]),
        ("AI-integrerte tiltak", [
            ["AI tidligvarsling", "Integrert AI-varsling på tvers av alle domener", "20 %", "15 %", "600 000"],
        ]),
    ]:
        b.heading(group, 4)
        b.table(
            ["Tiltak", "Beskrivelse", "P-red.", "Alv-red.", "Kostnad/år (NOK)"],
            rows,
            col_widths=[4.5, 7, 1.8, 2, 3],
        )
    b.note(
        "P-red. = reduksjon i hendelsessannsynlighet. Alv-red. = reduksjon i tap per hendelse. "
        "Tiltak kombineres multiplikativt (f.eks. 30 % + 20 % = 44 %, ikke 50 %)."
    )

    b.heading("3.5 Pooling-innstillinger", 2)
    b.body(
        "Aktivér med avkrysningsboksen **«Aktiver poolet PCC-scenario»** for å evaluere "
        "om operatøren drar nytte av å inngå i et risikodelingsbasseng."
    )
    b.heading("Bassengstruktur", 4)
    b.table(
        ["Felt", "Standard", "Beskrivelse"],
        [
            ["Antall medlemmer", "4", "Totalt antall operatører i bassenget inkl. denne. Min 2, maks 10."],
            ["Korrelasjonskoeffisient", "0,25", "Grad av samvariasjon mellom members' tapsår. 0 = uavhengige; 0,95 = nesten perfekt."],
            ["Likhetsvariasjonsbredde (±)", "0,15", "Hvor like de syntetiske bassengmedlemmene er. 0 = identiske; 0,50 = ±50 % variasjon."],
        ],
        col_widths=[5.5, 2.5, 10],
    )
    b.heading("Poolet gjenforsikring", 4)
    b.table(
        ["Felt", "Standard", "Beskrivelse"],
        [
            ["Egenrisiko for pool (Mill NOK)", "25", "Totalt behold for hele bassenget før GF trer inn."],
            ["GF-grense for pool (Mill NOK)", "400", "Maksimal gjenforsikringsdekning for hele bassenget."],
            ["GF-lastefaktor", "1,40", "GF-premie = forventet GF-utbetaling × lastefaktor."],
            ["Delt administrasjonsinsparing", "20 %", "Forventet kostnadsbesparelse på felles administrasjon."],
        ],
        col_widths=[5.5, 2.5, 10],
    )
    b.note(
        "v2.1: Bassengmedlemmene er syntetisk genererte – faktiske peer-data lastes ikke opp. "
        "Korrelasjonen er rangbasert blending (ikke Gaussisk kopula). "
        "Se Pooling-fanen i resultatpanelet for fullstendig liste over forenklinger."
    )

    # ── 4. Kjøreknapper ───────────────────────────────────────────────────────
    b.heading("4. Kjøreknapper og statuslinje", 1)
    b.table(
        ["Knapp", "Funksjon"],
        [
            ["Kjør analyse", "Sender inndata til backend, kjører Monte Carlo-simulering og returnerer resultater. Typisk kjøretid: 1–5 sekunder."],
            ["Last eksempel", "Fyller inn Nordic Aqua Partners AS (norsk lakseoppdretter, 9 200 t, 3 lokaliteter) som referansecase."],
            ["Nullstill", "Tømmer alle felt og resultater til standardverdier."],
        ],
        col_widths=[4, 14],
    )
    b.heading("Statuslinje – seks steg", 4)
    b.code_block(
        "① Bygg operatørmodell  →  ② Simuler tap  →  ③ Beregn strategier\n"
        "④ SCR-analyse          →  ⑤ Egnethetsvurdering  →  ⑥ Ferdig"
    )
    b.body("Aktivt steg er uthevet. Eventuelle feilmeldinger vises i rødt under statuslinjen.")

    # ── 5. Høyre panel ────────────────────────────────────────────────────────
    b.heading("5. Høyre panel – Resultater", 1)
    b.body("Resultater vises i åtte faner, tilgjengelige etter at en analyse er kjørt.")

    b.heading("5.1 Summary – Sammendrag", 2)
    b.body(
        "Gir et øyeblikksbilde av nøkkeltall for baseline-scenario og, dersom tiltak er valgt, "
        "for det mitigerte scenariet."
    )
    b.table(
        ["Nøkkeltall", "Beskrivelse"],
        [
            ["E[Årlig tap]", "Forventet gjennomsnittlig tap per år"],
            ["VaR 95 %", "Tap som overskrides i 5 % av årsscenarioene"],
            ["VaR 99,5 %", "Tap som overskrides i 0,5 % av årsscenarioene (≈ SCR)"],
            ["SCR (netto)", "Solvenskapitalkrav etter gjenforsikring og egenandelsstruktur"],
            ["Anbefalt strategi", "Verktøyets vurdering av best egnet risikostrategi"],
            ["Egnethetsscore", "Samlet score 0–100 for PCC-egnethet"],
            ["Konfidensnivå", "Grad av sikkerhet i anbefalingen (lav / moderat / høy)"],
        ],
        col_widths=[5, 12],
    )
    b.note(
        "Delta-kort (grønn = forbedring, rød = forverring) vises kun ved mitigert kjøring."
    )

    b.heading("5.2 Charts – Diagrammer", 2)
    b.body(
        "Viser opptil ti diagrammer generert av backend. "
        "Høyreklikk på et bilde og velg **Lagre bilde** for å laste det ned."
    )
    b.table(
        ["Diagram", "Innhold"],
        [
            ["Tapsdistribusjon", "Histogram over simulerte årstap med percentil-markører"],
            ["Kumulativ fordelingsfunksjon", "S-kurve; les av VaR direkte på y-aksen"],
            ["Domainefordeling", "Kakediagram: andel tap per risikodomene"],
            ["Strategi-TCOR", "Søylediagram: 5-års TCOR for alle fire strategier"],
            ["Kovariasjonsvarmekart", "Heatmap over korrelasjon mellom domener"],
            ["Scenariosammenligning", "Stablet søylediagram for p50/p95/p99 per domene"],
            ["Haledomenesammensetning", "Bidrag per domene i haleende av fordelingen"],
            ["Mitigeringseffekt", "Før/etter-sammenligning per tiltak (kun ved mitigert kjøring)"],
            ["Kostnads-/nyttekurve", "Tiltakskostnad vs. tapsnedgang per tiltak"],
        ],
        col_widths=[6, 11],
    )

    b.heading("5.3 Mitigation – Tiltak", 2)
    b.body("Detaljert analyse av de valgte tiltakenes effekt.")
    b.bullet("**Oversiktstabell**: Hvert valgt tiltak med domene, P-reduksjon, alvorlighetsreduksjon og estimert tapsnedgang")
    b.bullet("**Domenepåvirkning**: Prosentvis reduksjon per risikodomene")
    b.bullet("**Tapsnedgang (NOK)**: Estimert reduksjon i forventet årlig tap")
    b.bullet("**SCR-reduksjon (NOK)**: Estimert reduksjon i solvenskapitalkrav")

    b.heading("5.4 Recommendation – Anbefaling", 2)
    b.body("Kjernevurderingen av PCC-egnethet basert på seks kriterier.")
    b.heading("Vurderingsbanner", 4)
    b.table(
        ["Utfall", "Betingelse"],
        [
            ["✓  ANBEFALT", "Alle seks kriterier bestått; composite score ≥ 70"],
            ["⚠  POTENSIELT EGNET", "Minst fire kriterier bestått; moderate kostnadsforhold"],
            ["✗  IKKE ANBEFALT", "Kritisk kriterium feilet (f.eks. PCC-kostnad > 15 % over FI-kostnad)"],
        ],
        col_widths=[5, 12],
    )
    b.heading("Seks kriterier", 4)
    b.table(
        ["Kriterium", "Hva måles"],
        [
            ["Tapsstabilitet", "CV av simulerte tap – lav CV gir mer forutsigbar finansiering i PCC"],
            ["Kapitaleffektivitet", "Forholdet mellom SCR og egenkapital"],
            ["Premieeffektivitet", "PCC-totalkostnad versus markedspremie"],
            ["Selvforsikringsegnethet", "Om operatøren kan finansiere worst-case tap"],
            ["Diversifikasjonspotensial", "Mulighet for å spre risiko mellom lokaliteter"],
            ["Tapskonsentrasjon", "Grad av dominans fra ett enkelt risikodomene"],
        ],
        col_widths=[5, 12],
    )

    b.heading("5.5 Allocation – Eksponeringsfordeling", 2)
    b.body(
        "Viser hvordan operatørinndataene er oversatt til den fullstendige interne modellen."
    )
    b.table(
        ["Nøkkeltall", "Beskrivelse"],
        [
            ["Biomasse-TIV-ratio", "Forholdet mellom oppgitt biomasseverdi og malobjektets totale forsikringsverdi"],
            ["Antatt TIV", "Total forsikringsverdi som brukes i modellen"],
            ["SCR / Egenkapital", "Solvens-ratio; bør være < 40 % for at PCC skal være kapitalmessig forsvarlig"],
            ["E[tap] / EBITDA", "Tapsbyrde relativt til driftsresultat"],
            ["VaR 99,5 % / FCF", "Worst-case tap versus fri kontantstrøm"],
        ],
        col_widths=[5, 12],
    )
    b.note(
        "Et oransje advarselsbanner vises dersom anvendt biomasseverdi avviker mer enn 10 % "
        "fra den foreslåtte verdien."
    )

    b.heading("5.6 Loss History – Taphistorikk", 2)
    b.body("Vises kun dersom operatørmalen inneholder historiske tapsposter.")
    b.table(
        ["Hendelsestype", "Domene"],
        [
            ["Mortalitet, sykdom, HAB, lus", "Biologisk"],
            ["Utstyrsskade, not, fortøyning, storm", "Strukturell"],
            ["Driftsavbrudd, tyveri, personell", "Operasjonell"],
            ["Andre / ukjente typer", "Ukjent"],
        ],
        col_widths=[8, 9],
    )
    b.note(
        "Gjennomsnittlig årstap per domene beregnes over hele observasjonsperioden, "
        "ikke bare de år domenet har hendelser."
    )

    b.heading("5.7 Pooling – Sammenslutningsanalyse", 2)
    b.body("Vises kun dersom Pooling er aktivert og analyse er kjørt.")
    b.table(
        ["Seksjon", "Innhold"],
        [
            ["A. Vurderingsbanner", "Viser LEVEDYKTIG (grønn) dersom pooled TCOR < standalone TCOR og CV-reduksjon > 5 %"],
            ["B. Diversifikasjonsgevinst", "Standalone CV vs. Pooled CV, VaR-reduksjon med %-badge"],
            ["C. Økonomisammenligning", "Tabell: GF-premie, SCR og 5-års TCOR – standalone vs. poolet"],
            ["D. Bassengmedlemsoversikt", "Per-medlemstabell; operatøren fremhevet med blå bakgrunn"],
            ["E. Modellforutsetninger", "Nedtrekkbar liste over v2.1-forenklinger (viktig for styrekommunikasjon)"],
        ],
        col_widths=[5, 12],
    )

    b.heading("5.8 Report – PDF-rapport", 2)
    b.body(
        "Viser nedlastingsstatus for den genererte PDF-rapporten. "
        "Klikk på lenken for å åpne rapporten i nettleseren. "
        "Rapporten er et 13-siders styredokument med sammendrag, Monte Carlo-resultater, "
        "domenekorrelert risikoanalyse, scenariosammenligning og PCC-kostnadsstruktur."
    )

    # ── 6. Arbeidsflyt ────────────────────────────────────────────────────────
    b.heading("6. Typisk arbeidsflyt", 1)
    b.heading("Enkel vurdering (ca. 10 minutter)", 3)
    b.table(
        ["Steg", "Handling"],
        [
            ["1", "Klikk **Last eksempel** – fyller inn Nordic Aqua-profil"],
            ["2", "Klikk **Kjør analyse** – venter ~2 sekunder"],
            ["3", "Les **Recommendation**-fanen – er PCC anbefalt?"],
            ["4", "Åpne **Summary**-fanen – sammenlign TCOR for alle strategier"],
            ["5", "Last ned **PDF-rapport** – del med styre og rådgiver"],
        ],
        col_widths=[1.5, 16],
    )
    b.heading("Fullstendig vurdering med tiltak (ca. 20 minutter)", 3)
    b.table(
        ["Steg", "Handling"],
        [
            ["1", "Fyll inn Operatørprofil med faktiske tall"],
            ["2", "Aktiver relevante risikoreduserende tiltak i Mitigation-seksjonen"],
            ["3", "Kjør analyse"],
            ["4", "Les **Mitigation**-fanen – hvilke tiltak gir størst verdi?"],
            ["5", "Sammenlign baseline og mitigert i **Summary**"],
            ["6", "Last ned PDF"],
        ],
        col_widths=[1.5, 16],
    )
    b.heading("Pooling-evaluering", 3)
    b.table(
        ["Steg", "Handling"],
        [
            ["1", "Fullfør standard vurdering (steg 1–4 ovenfor)"],
            ["2", "Åpne **Pooling**-seksjonen i venstrekolonnen og aktiver pooling"],
            ["3", "Juster evt. antall medlemmer og korrelasjonskoeffisient"],
            ["4", "Kjør analyse på nytt"],
            ["5", "Les **Pooling**-fanen – er pooled PCC levedyktig?"],
            ["6", "Sammenlign TCOR: standalone PCC vs. Pooled PCC vs. Full forsikring"],
        ],
        col_widths=[1.5, 16],
    )

    # ── 7. Feltforklaringer ───────────────────────────────────────────────────
    b.heading("7. Feltforklaringer – oppslagstabell", 1)
    b.table(
        ["Felt / parameter", "Norsk forklaring", "Enhet"],
        [
            ["`n_sites`", "Antall produksjonslokaliteter", "stk"],
            ["`total_biomass_tonnes`", "Total stående biomasse", "tonn"],
            ["`biomass_value_per_tonne`", "Forsikringsverdi per tonn biomasse", "NOK/t"],
            ["`annual_revenue_nok`", "Estimert årsomsetning", "NOK"],
            ["`annual_premium_nok`", "Estimert forsikringspremie", "NOK"],
            ["`n_simulations`", "Antall Monte Carlo-iterasjoner", "stk"],
            ["`domain_correlation`", "Korrelasjonsprofil for risikodomener", "kode"],
            ["`generate_pdf`", "Generer PDF-styrerapport", "ja/nei"],
            ["`use_history_calibration`", "Kalibrér mot historiske tap", "ja/nei"],
            ["`retention_nok`", "Egenrisiko-grense", "NOK"],
            ["`pooled_retention_nok`", "Total basseng-egenrisiko", "NOK"],
            ["`pooled_ri_limit_nok`", "Maksimal basseng-GF-dekning", "NOK"],
            ["`pooled_ri_loading_factor`", "GF-premie = E[utbetaling] × faktor", "×"],
            ["`inter_member_correlation`", "Korrelasjon mellom bassengmedlemmer", "0–1"],
            ["`similarity_spread`", "Variasjon i syntetiske jevnaldrende", "± fraksjon"],
            ["`shared_admin_saving_pct`", "Forventet administrasjonsbesparelse", "%"],
            ["`allocation_basis`", "Kostnadsforderingsgrunnlag i bassenget", "kode"],
        ],
        col_widths=[5, 9, 4],
    )

    # ── 8. Tolkning ───────────────────────────────────────────────────────────
    b.heading("8. Tolkning av nøkkeltall", 1)
    b.heading("Relative tapsnivåer", 3)
    b.table(
        ["Indikator", "Lavt", "Moderat", "Høyt"],
        [
            ["E[tap] / Biomasse-TIV", "< 1 %", "1–3 %", "> 3 %"],
            ["CV", "< 1,5", "1,5–2,5", "> 2,5"],
            ["VaR 99,5 % / E[tap]", "< 5×", "5–8×", "> 8×"],
        ],
        col_widths=[6, 3.5, 3.5, 3.5],
    )
    b.heading("SCR-vurdering", 3)
    b.table(
        ["SCR-nivå", "Tolkning"],
        [
            ["< 10 M NOK", "PCC kan etableres med minimal kapital. Svært gunstig."],
            ["10–50 M NOK", "Typisk for mellomstore operatører. PCC er gjerne levedyktig med riktig GF-struktur."],
            ["> 100 M NOK", "Krever betydelig egenkapitalreserve. Pooling eller hybrid kan være bedre."],
        ],
        col_widths=[4, 14],
    )
    b.heading("Egnethetsscore", 3)
    b.table(
        ["Score", "Tolkning"],
        [
            ["80–100", "PCC er klart anbefalt"],
            ["60–79", "PCC kan fungere; nøye vurdering anbefales"],
            ["40–59", "Betinget egnet; vurder tiltak eller hybrid"],
            ["< 40", "PCC anbefales ikke i nåværende konfigurasjon"],
        ],
        col_widths=[3, 15],
    )
    b.heading("Poolingfordel", 3)
    b.table(
        ["CV-reduksjon", "Tolkning"],
        [
            ["> 20 %", "Sterk diversifikasjonsgevinst – pooling klart fordelaktig"],
            ["10–20 %", "Moderat gevinst – pooling bør vurderes"],
            ["5–10 %", "Marginal gevinst – avhenger av kostnadsstruktur"],
            ["< 5 %", "Liten diversifikasjonsgevinst – pooling gir begrenset fordel"],
        ],
        col_widths=[4, 14],
    )

    # ── 9. Feilsøking ─────────────────────────────────────────────────────────
    b.heading("9. Feilsøking", 1)
    b.table(
        ["Feilmelding / symptom", "Løsning"],
        [
            ["«Failed to connect to backend»",
             "Backend kjører ikke. Kjør: uvicorn backend.main:app --reload --port 8000"],
            ["«Failed to load example»",
             "Sjekk at backend er oppe og at data/sample_input.json eksisterer."],
            ["Analysen henger – spinner ferdig aldri",
             "Sjekk backend-loggene. Prøv n_simulations = 1 000 for rask test."],
            ["PDF-lenken fungerer ikke",
             "Kontroller at mappen backend/static/reports/ eksisterer og at backend har skriverettigheter."],
            ["Pooled CV ikke lavere enn standalone CV",
             "Kan skje ved svært høy korrelasjonskoeffisient (> 0,7) eller kun 2 medlemmer. Prøv lavere korrelasjon eller flere medlemmer."],
            ["Kalibreringsstatus viser «Ikke aktiv»",
             "Historikkalibrering må aktiveres eksplisitt med bryteren «Bruk historikkalibrering» i Modellinnstillinger."],
        ],
        col_widths=[6, 12],
    )

    # ── 10. Arkitektur ────────────────────────────────────────────────────────
    b.heading("10. Teknisk arkitektur (kortfattet)", 1)
    b.code_block(
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│  Nettleser  (localhost:5173)                                │\n"
        "│  React 18 + Vite 5                                          │\n"
        "│  ├── Venstre panel: accordion-inndataskjema                 │\n"
        "│  └── Høyre panel:   resultattabs (8 faner)                  │\n"
        "└──────────────────────┬──────────────────────────────────────┘\n"
        "                       │  POST /api/feasibility/run  (JSON)\n"
        "                       ▼\n"
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│  FastAPI backend  (localhost:8000)                          │\n"
        "│  ├── backend/api/feasibility.py   – HTTP-adapter            │\n"
        "│  ├── backend/services/operator_builder.py                   │\n"
        "│  └── backend/services/run_analysis.py  – full pipeline      │\n"
        "└──────────────────────┬──────────────────────────────────────┘\n"
        "                       │  Python-modulkall\n"
        "                       ▼\n"
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│  Analysemotor (eksisterende Python-moduler)                 │\n"
        "│  ├── models/monte_carlo.py          – simulering            │\n"
        "│  ├── models/domain_correlation.py   – domenematrise        │\n"
        "│  ├── models/strategies/             – fire strategier       │\n"
        "│  ├── models/pooling/                – bassenganalyse        │\n"
        "│  ├── analysis/mitigation.py         – 12 tiltak             │\n"
        "│  ├── analysis/suitability_engine.py – 6-kriterie-score      │\n"
        "│  ├── reporting/chart_generator.py   – 10 diagrammer         │\n"
        "│  └── reporting/pdf_report.py        – 13-siders PDF         │\n"
        "└─────────────────────────────────────────────────────────────┘"
    )

    b.heading("Dataflyt", 4)
    steps = [
        "Brukeren fyller inn skjemaet og klikker **Kjør analyse**",
        "Frontend sender én POST-forespørsel til `/api/feasibility/run`",
        "Backend bygger en fullstendig `OperatorInput` fra den forenklede profilen",
        "Monte Carlo-motoren kjører N simuleringer (vektorisert NumPy, ~1 sek for 5 000)",
        "Fire (eller fem) strategier beregnes",
        "Egnethetsmotoren skårer 6 kriterier og setter en anbefaling",
        "Valgfritt: PDF genereres og lagres i `backend/static/reports/`",
        "Hele svaret returneres som JSON; frontend rendrer resultattabsene",
    ]
    for i, step in enumerate(steps, 1):
        b.bullet(f"**{i}.** {step}")

    # Footer
    doc.add_paragraph()
    b.horizontal_rule()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "Shield Risk Platform v2.1  —  © Shield Risk Consulting  "
        "—  Konfidensielt. Ikke distribuer uten tillatelse."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.italic = True


if __name__ == "__main__":
    out_path = Path(__file__).resolve().parent.parent / "BRUKERMANUAL_GUI.docx"
    b = DocBuilder()
    build(b)
    b.doc.save(out_path)
    print(f"Saved: {out_path}")
